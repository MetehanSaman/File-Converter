import os
import customtkinter
from customtkinter import filedialog
from tkinter import *
import img2pdf
from docx2pdf import convert
from pdf2docx import Converter
from docx import Document
from PyPDF2 import PdfMerger
import pypdfium2 as pdfium
import sys


root = customtkinter.CTk()
root.geometry("650x500")
root.title("File Converter")

customtkinter.set_appearance_mode("Dark")
customtkinter.set_default_color_theme("blue")
switch = True

options = ["PDF to JPG", "PDF to DOC", "DOC to JPG", "DOC to PDF", "JPG to PDF"]
i = 0
k = 0
output = ""
outputdir = ""
Desktop = ""
path_ = ""
itemname = ""
picked = StringVar()
picked.set(options[0])
function = "PDF to JPG"
itemnames = []
extension = ".pdf"
check_switch = 0
check_switch1 = 0
check_switch2 = 0
check = customtkinter.StringVar()
check1 = customtkinter.StringVar()
check2 = customtkinter.StringVar()
merge = 0
merge1 = 0
merge2 = 0
merger = PdfMerger()
show_check = 0
show_check1 = 0
show_check2 = 0


def show_button():
    global checkbox
    global check_switch
    global show_check
    if show_check == 0:
        checkbox = customtkinter.CTkCheckBox(master=sag_ust_frame, text="Merge JPGs", variable=check, command=Check)
        checkbox.place(relx=0.05, rely=0.7, relwidth=0.90, relheight=0.25)
        check_switch = 1
        show_check = 1
    elif show_check == 1:
        pass


def show_button1():
    global checkbox1
    global check_switch1
    global show_check1

    if show_check1 == 0:
        checkbox1 = customtkinter.CTkCheckBox(master=sag_ust_frame, text="Merge PDFs", variable=check1, command=Check1)
        checkbox1.place(relx=0.05, rely=0.7, relwidth=0.90, relheight=0.25)
        check_switch1 = 1
        show_check1 = 1
    elif show_check1 == 1:
        pass


def show_button2():
    global checkbox2
    global check_switch2
    global show_check2

    if show_check2 == 0:
       checkbox2 = customtkinter.CTkCheckBox(master=sag_ust_frame, text="Merge DOCs", variable=check2, command=Check2)
       checkbox2.place(relx=0.05, rely=0.7, relwidth=0.90, relheight=0.25)
       check_switch2 = 1
       show_check2 = 1
    elif show_check2 == 1:
        pass


def hide_button(widget):
    global check_switch
    global show_check

    show_check = 0
    widget.destroy()
    check_switch = 0


def hide_button1(widget):
    global check_switch1
    global show_check1

    show_check1 = 0
    widget.destroy()
    check_switch1 = 0


def hide_button2(widget):
    global check_switch2
    global show_check2

    show_check2 = 0
    widget.destroy()
    check_switch2 = 0


def checkbox_check():
    global check_switch
    if check_switch == 0:
        pass
    elif check_switch == 1:
        hide_button(checkbox)


def checkbox_check1():
    global check_switch1
    if check_switch1 == 0:
        pass
    elif check_switch1 == 1:
        hide_button1(checkbox1)


def checkbox_check2():
    global check_switch2
    if check_switch2 == 0:
        pass
    elif check_switch2 == 1:
        hide_button2(checkbox2)


def selected(variable):
    global function
    global extension
    global itemnames
    for a in itemnames:
        listbox.delete(ACTIVE)
    itemnames = []
    function = picked.get()
    match function:
        case "PDF to JPG":
            if path_ == "":
                convert_label.configure(text="Please Choose a Path")
            else:
                convert_label.configure(text="Converting Bekleniyor")
                checkbox_check()
                checkbox_check1()
                checkbox_check2()
                extension = ".pdf"
                updlistbox()
        case "PDF to DOC":
            if path_ == "":
                convert_label.configure(text="Please Choose a Path")
            else:
                convert_label.configure(text="Converting Bekleniyor")
                show_button1()
                extension = ".pdf"
                updlistbox()
        case "DOC to JPG":
            if path_ == "":
                convert_label.configure(text="Please Choose a Path")
            else:
                convert_label.configure(text="Converting Bekleniyor")
                checkbox_check()
                checkbox_check1()
                checkbox_check2()
                extension = ".docx"
                updlistbox()
        case "DOC to PDF":
            if path_ == "":
                convert_label.configure(text="Please Choose a Path")
            else:
                convert_label.configure(text="Converting Bekleniyor")
                show_button2()
                extension = ".docx"
                updlistbox()
        case "JPG to PDF":
            if path_ == "":
                convert_label.configure(text="Please Choose a Path")
            else:
                convert_label.configure(text="Converting Bekleniyor")
                show_button()
                extension = ".jpg"
                updlistbox()


def updlistbox():
    dirname = path_
    for fname in os.listdir(dirname):
        if not fname.endswith(extension):
            continue
        filename = os.path.join(fname)
        if os.path.isdir(filename):
            continue
        elif fname in itemnames:
            itemnames.append(filename)
        listbox.insert(END, fname)


def pickpath():
    global path_
    global itemname
    global extension
    global itemnames

    for a in itemnames:
        listbox.delete(ACTIVE)

    path_ = filedialog.askdirectory(initialdir="/", title="Select a File")
    path_ = path_.replace(os.sep, '/')
    itemname = path_.split('/')[-1]
    if path_ == "":
        convert_label.configure(text="Please Choose a Path")
    else:
        convert_label.configure(text="Converting Bekleniyor")
        updlistbox()


def drop():
    global itemnames
    pick = listbox.get(ANCHOR)
    if pick == "":
        convert_label.configure(text="Please Choose a File")
    else:
        convert_label.configure(text="Converting Bekleniyor")
        itemnames.remove(pick)
        listbox.delete(ANCHOR)


def appearance():
    global switch
    if not switch:
        customtkinter.set_appearance_mode("Dark")
        switch = True
    else:
        customtkinter.set_appearance_mode("Light")
        switch = False


def delete_log():
    global outputdir

    sys.stderr.close()
    os.remove(os.path.join('{}/'.format(outputdir), "log.txt"))

def start():
    global function
    global itemnames

    if path_ == "":
        convert_label.configure(text="Please Choose a Path")
    else:
        convert_label.configure(text="Converting Bekleniyor")
        if itemnames == []:
            convert_label.configure(text="Files Not Found")
        else:
            match function:
                case "PDF to JPG":
                    pdf2img()
                    convert_label.configure(text="Converting Tamamlandı")
                    delete_log()
                case "PDF to DOC":
                    pdf2doc()
                    convert_label.configure(text="Converting Tamamlandı")
                    delete_log()
                case "DOC to JPG":
                    doc2img()
                    convert_label.configure(text="Converting Tamamlandı")
                    delete_log()
                case "DOC to PDF":
                    doc2pdf()
                    convert_label.configure(text="Converting Tamamlandı")
                    delete_log()
                case "JPG to PDF":
                    image2pdf()
                    convert_label.configure(text="Converting Tamamlandı")
                    delete_log()


def Check():
    global check
    global merge
    if check.get() == "1":
        merge = 1

    elif check.get() == "0":
        merge = 0

def Check1():
    global check1
    global merge1
    if check1.get() == "1":
        merge1 = 1

    elif check1.get() == "0":
        merge1 = 0

def Check2():
    global check2
    global merge2
    if check2.get() == "1":
        merge2 = 1

    elif check2.get() == "0":
        merge2 = 0



sol_ust_frame = customtkinter.CTkFrame(master=root,
                                          width=500,
                                          height=500,
                                          corner_radius=10)

sol_ust_frame.place(relx=0.0266, rely=0.04, relwidth=0.66666, relheight=0.2)



Choosebutton = customtkinter.CTkButton(master=sol_ust_frame, text="Choose Folder Path ", command=pickpath)
Choosebutton.place(relx=0.05, rely=0.15, relwidth=0.90, relheight=0.70)

sol_alt_frame = customtkinter.CTkFrame(master=root,
                                          width=500,
                                          height=500,
                                          corner_radius=10)

sol_alt_frame.place(relx=0.0266, rely=0.25, relwidth=0.66666, relheight=0.71)

listbox = Listbox(master=sol_alt_frame, width=500, height=500)

listbox.place(relx=0.05, rely=0.05, relwidth=0.90, relheight=0.65)

for i in itemnames:
    listbox.insert(END, i)

drop_button = customtkinter.CTkButton(sol_alt_frame, text="Drop", command=drop)
drop_button.place(relx=0.3, rely=0.75, relwidth=0.4, relheight=0.15)



sag_ust_frame = customtkinter.CTkFrame(master=root,
                                          width=500,
                                          height=500,
                                          corner_radius=10)
sag_ust_frame.place(relx=0.6996, rely=0.04, relwidth=0.26666, relheight=0.20)

combobox = customtkinter.CTkOptionMenu(master=sag_ust_frame,
                                       values=options,
                                       command=selected,
                                       variable=picked)

combobox.place(relx=0.05, rely=0.35, relwidth=0.90, relheight=0.30)



sag_orta_frame = customtkinter.CTkFrame(master=root,
                                          width=500,
                                          height=500,
                                          corner_radius=10)
sag_orta_frame.place(relx=0.6996, rely=0.25, relwidth=0.26666, relheight=0.50)

sag_alt_frame = customtkinter.CTkFrame(master=root,
                                          width=500,
                                          height=500,
                                          corner_radius=10)
sag_alt_frame.place(relx=0.6996, rely=0.76, relwidth=0.26666, relheight=0.20)

convert_label = customtkinter.CTkLabel(master=sag_alt_frame, width=15, text="Converting Bekleniyor")
convert_label.place(relx=0.03, rely=0.3, relwidth=0.94, relheight=0.4)

dark_light_switch = customtkinter.CTkSwitch(sag_orta_frame, text=" Dark/Light", command=lambda: appearance())
dark_light_switch.place(relx=0.15, rely=0.65, relwidth=0.70, relheight=0.20)


start_button = customtkinter.CTkButton(sag_orta_frame, text=" Start", command=start)
start_button.place(relx=0.125, rely=0.15, relwidth=0.75, relheight=0.25)


def desktop():
    global Desktop
    global outputdir

    Desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
    Desktop = Desktop.replace(os.sep, '/')
    if 'ConvertedItems' in os.listdir('{}'.format(Desktop)):
        pass
    else:
        os.mkdir('{}/ConvertedItems'.format(Desktop))
    outputdir = '{}/ConvertedItems'.format(Desktop)

    sys.stderr = open("{}/log.txt".format(outputdir), "w")


def converteditem():
    global Desktop
    global output

    if itemname in os.listdir('{}'.format(outputdir)):
        pass
    else:
        os.mkdir('{}/ConvertedItems/{}'.format(Desktop, itemname))
    output = '{}/ConvertedItems/{}'.format(Desktop, itemname)


def pdf2img():
    global Desktop
    global i
    global itemnames
    dirname = path_
    desktop()
    imgs = []
    converteditem()

    for fname in os.listdir(dirname):
        if not fname.endswith(".pdf"):
            continue
        if fname in itemnames:
            path = os.path.join(dirname, fname)
            if os.path.isdir(path):
                continue
            imgs.append(path),

            pdf = pdfium.PdfDocument(imgs[i])
            n_pages = len(pdf)
            fname_ = fname.split(".")[0]

            for page_number in range(n_pages):
                page = pdf.get_page(page_number)
                pil_image = page.render(scale=1, rotation=0, crop=(0, 0, 0, 0)).to_pil()
                pil_image.save(output + "/" + fname_ + "_" + "{}.jpg".format(page_number+1))
            i = i + 1

    i = 0
    return


def image2pdf():
    dirname = path_
    imgs = []
    desktop()
    converteditem()
    name_item = itemname
    if merge == 0:
        for fname in os.listdir(dirname):
            if not fname.endswith(".jpg"):
                continue
            if fname in itemnames:
                path = os.path.join(dirname, fname)
                itemname_ = fname.split('.')[0]
                if os.path.isdir(path):
                    continue
                imgs.append(path)
                with open(output + "/" + itemname_ + ".pdf", "wb") as f:
                    f.write(img2pdf.convert(imgs)) # type: ignore
                imgs = []

    elif merge == 1:

        name_check = 1

        for fname in os.listdir(dirname):
            if not fname.endswith(".jpg"):
                continue
            if fname in itemnames:
                path = os.path.join(dirname, fname)
                if os.path.isdir(path):
                    continue
                imgs.append(path)
        while True:
            if name_item + ".pdf" in os.listdir(output):
                name_item = itemname + "-" + "{}".format(name_check)
                name_check = name_check + 1
                continue
            else:
                with open(output + "/" + name_item + ".pdf", "wb") as f:
                    f.write(img2pdf.convert(imgs)) # type: ignore
                break
    return


def doc2pdf():
    global merger
    desktop()
    converteditem()
    dirname = path_
    imgs = []
    m = 0
    name_item = itemname

    if merge2 == 0:
        for fname in os.listdir(dirname):
            if not fname.endswith(".docx" or "doc"):
                continue
            if fname in itemnames:
                path = os.path.join(dirname, fname)
                if os.path.isdir(path):
                    continue
                imgs.append(path)
                convert(imgs[m], output)
                m = m + 1

    elif merge2 == 1:

        name_check = 1

        if 'Temp' in os.listdir('{}'.format(outputdir)):
            pass
        else:
            os.mkdir('{}/Temp/'.format(outputdir))

        for fname in os.listdir(dirname):
            if not fname.endswith(".docx" or "doc"):
                continue
            if fname in itemnames:
                path = os.path.join(dirname, fname)
                itemname_ = fname.split('.')[0]
                if os.path.isdir(path):
                    continue
                imgs.append(path)
                convert(imgs[m], '{}/Temp/'.format(outputdir))
                m = m+1
                merger.append('{}/Temp/'.format(outputdir) + itemname_ + ".pdf")

        while True:
            if name_item + ".pdf" in os.listdir(output):
                name_item = itemname + "-" + "{}".format(name_check)
                name_check = name_check + 1
                continue
            else:
                merger.write(output + "/" + name_item + ".pdf")
                break

        merger.close()
        merger = PdfMerger()
        for f in os.listdir('{}/Temp/'.format(outputdir)):
            os.remove(os.path.join('{}/Temp/'.format(outputdir), f))
        os.rmdir('{}/Temp/'.format(outputdir))

    return


def pdf2doc():
    global k
    global itemname
    name_item = itemname
    dirname = path_
    imgs = []
    desktop()
    converteditem()
    temp = []
    if merge1 == 0:
        for fname in os.listdir(dirname):
            if not fname.endswith(".pdf"):
                continue
            if fname in itemnames:
                path = os.path.join(dirname, fname)
                itemname_ = fname.split('.')[0]
                if os.path.isdir(path):
                    continue
                imgs.append(path),
                cv = Converter(imgs[k])
                cv.convert(output + "/" + itemname_ + ".docx")
                cv.close()
                k = k + 1

    elif merge1 == 1:

        name_check = 1

        if 'Temp' in os.listdir('{}'.format(outputdir)):
            pass
        else:
            os.mkdir('{}/Temp/'.format(outputdir))

        for fname in os.listdir(dirname):
            if not fname.endswith(".pdf"):
                continue
            if fname in itemnames:
                path = os.path.join(dirname, fname)
                itemname_ = fname.split('.')[0]
                if os.path.isdir(path):
                    continue
                imgs.append(path),
                cv = Converter(imgs[k])
                cv.convert(outputdir + "/Temp/" + itemname_ + ".docx")
                temp.append(outputdir + "/Temp/" + itemname_ + ".docx")
                cv.close()
                k = k + 1

        merged_document = Document()

        for index, file in enumerate(temp):
            sub_doc = Document(file)

            if index < len(temp) - 1:
                sub_doc.add_page_break()

            for element in sub_doc.element.body:
                merged_document.element.body.append(element)

        while True:
            if name_item + ".docx" in os.listdir(output):
                name_item = itemname + "-" + "{}".format(name_check)
                name_check = name_check + 1
                continue
            else:
                merged_document.save(output + "/" + name_item + ".docx")
                break

        for f in os.listdir(outputdir + "/Temp/"):
            os.remove(os.path.join(outputdir + "/Temp/", f))
        os.rmdir(outputdir + "/Temp/")

    k = 0
    return


def doc2img():
    global path_
    global Desktop
    global i
    global itemnames
    imgs = []
    oldpath_ = path_
    desktop()
    converteditem()
    global itemnames
    olditemnames = itemnames

    if 'Temp_doc' in os.listdir('{}'.format(outputdir)):
        pass
    else:
        os.mkdir('{}/Temp_doc/'.format(outputdir))

    output_ = '{}/Temp_doc/'.format(outputdir)
    for l in itemnames:
        path_ = path_ + "/" + l
        convert(path_, output_)
        path_ = oldpath_
    path_ = output_
    dirname = output_
    itemnames = []
    for fname in os.listdir(dirname):
        if not fname.endswith(".pdf"):
            continue
        filename = os.path.join(fname)
        if os.path.isdir(filename):
            continue
        if not fname in itemnames:
            itemnames.append(filename)

    for fname in os.listdir(dirname):
        if not fname.endswith(".pdf"):
            continue
        if fname in itemnames:
            path = os.path.join(dirname, fname)
            if os.path.isdir(path):
                continue
            imgs.append(path),

            pdf = pdfium.PdfDocument(imgs[i])
            n_pages = len(pdf)
            fname_ = fname.split(".")[0]

            for page_number in range(n_pages):
                page = pdf.get_page(page_number)
                pil_image = page.render(scale=1, rotation=0, crop=(0, 0, 0, 0)).to_pil()
                pil_image.save(output + "/" + fname_ + "_" + "{}.jpg".format(page_number+1))

            pdf.close()
            i = i + 1
    i = 0

    itemnames = olditemnames
    path_ = oldpath_
    for f in os.listdir(output_):
        os.remove(os.path.join(output_, f))
    os.rmdir(output_)

    return


root.mainloop()
